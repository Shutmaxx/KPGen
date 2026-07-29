# -*- coding: utf-8 -*-
"""Утилиты для правки текста в DOCX.

Word разбивает текст абзаца на несколько фрагментов (runs) — из-за проверки
орфографии, правок и смены начертания. Поэтому наивный поиск подстроки внутри
одного run не находит текст, который визуально выглядит цельным.
Здесь замена делается по склеенному тексту абзаца с сохранением форматирования.
"""
from __future__ import annotations

from typing import Iterable

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def iter_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    """Все абзацы документа: тело, таблицы, колонтитулы."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    """Заменяет подстроку в абзаце, сохраняя оформление.

    Сначала подстановка выполняется внутри отдельных фрагментов (runs) —
    так сохраняется разное начертание в одном абзаце, например жирное
    название компании и обычный текст после него. Склейка всего абзаца
    применяется только тогда, когда плейсхолдер разорван между фрагментами.

    Возвращает True, если замена произошла.
    """
    runs = paragraph.runs
    if not runs:
        return False

    # Обычный случай: плейсхолдер целиком лежит в одном фрагменте.
    replaced_inside = False
    for run in runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            replaced_inside = True
    if replaced_inside:
        return True

    # Запасной путь: Word разорвал плейсхолдер между фрагментами.
    full = "".join(run.text for run in runs)
    if old not in full:
        return False

    runs[0].text = full.replace(old, new)
    for run in runs[1:]:
        run.text = ""
    return True


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Полностью заменяет текст абзаца, сохраняя формат первого run.

    Гиперссылки удаляются: их текст лежит вне runs абзаца и иначе
    остался бы в документе рядом с новым значением.
    """
    for hyperlink in paragraph._p.findall(qn("w:hyperlink")):
        paragraph._p.remove(hyperlink)

    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def replace_everywhere(doc: DocxDocument, mapping: dict[str, str]) -> dict[str, int]:
    """Подставляет значения плейсхолдеров по всему документу.

    Возвращает счётчик замен по каждому ключу — по нему видно,
    какие плейсхолдеры в шаблоне не найдены.
    """
    counts = {key: 0 for key in mapping}
    for paragraph in iter_paragraphs(doc):
        for key, value in mapping.items():
            if replace_in_paragraph(paragraph, key, value):
                counts[key] += 1
    return counts
