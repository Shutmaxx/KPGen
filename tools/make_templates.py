# -*- coding: utf-8 -*-
"""Готовит шаблоны с плейсхолдерами из рабочих файлов менеджера.

Запускается один раз (и повторно, если менеджер изменит исходные файлы).
Оформление, бланк, логотипы и стили остаются нетронутыми — меняется
только текст в тех местах, которые зависят от клиента и разговора.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

from app.services.docx_utils import set_paragraph_text, replace_in_paragraph

ROOT = Path(__file__).resolve().parents[2]          # D:\Петров
PKG = Path(__file__).resolve().parents[1]           # kpgen
TEMPLATES = PKG / "templates"

SRC_KP = ROOT / "КП_Фанагория_обновленное.docx"
SRC_PRES = ROOT / "Презентация_для_Акрон_Холдинг.pptx"
DST_KP = TEMPLATES / "kp_template.docx"
DST_PRES = TEMPLATES / "presentation_template.pptx"

# Абзац -> плейсхолдер. Индексы выверены по исходному файлу.
# Абзац 2 собирается особо: краткое название выделяется жирным,
# остальной текст — обычным начертанием.
PARAGRAPH_PLACEHOLDERS = {
    3: "{{ПОТРЕБНОСТИ_КЛИЕНТА}}",
    18: "{{ЧТО_ПРЕДЛАГАЕМ}}",
    20: "{{МЕНЕДЖЕР_ФИО}}",
    21: "{{МЕНЕДЖЕР_ДОЛЖНОСТЬ}}",
    22: "{{МЕНЕДЖЕР_ТЕЛЕФОН_1}}",
    23: "{{МЕНЕДЖЕР_ТЕЛЕФОН_2}}",
    24: "{{МЕНЕДЖЕР_ТЕЛЕФОН_3}}",
    25: "{{МЕНЕДЖЕР_EMAIL}}",
    26: "{{МЕНЕДЖЕР_САЙТ}}",
}

# Абзац с описанием компании: жирное название + обычный текст.
COMPANY_PARAGRAPH = 2
COMPANY_BOLD_PLACEHOLDER = "{{КОМПАНИЯ_КРАТКО}}"
# Разделитель не зашит в шаблон: он зависит от того, начинается ли
# характеристика со слова «является» или с описания без глагола.
COMPANY_TAIL_PLACEHOLDER = "{{ПОЛЬЗА_В_СЕГМЕНТЕ}}"

# Точечные замены упоминаний конкретного клиента в статичных блоках.
INLINE_REPLACEMENTS = [
    (
        "Бренд и безопасность. Мы не заставляем Ваших контрагентов искать Вас "
        "в общем каталоге. Мы внедрим персональный i-frame модуль прямо "
        "на корпоративный портал «Фанагории».",
        "Бренд и безопасность. Контрагенты могут найти Ваши закупки не только "
        "на нашей площадке и агрегаторах, но и на Вашем сайте. Мы можем "
        "разработать персональный i-frame модуль для внедрения прямо "
        "на корпоративный портал {{КОМПАНИЯ_КРАТКО}}.",
    ),
]

# Замена текста готовых блоков по требованию менеджера.
BLOCK_REPLACEMENTS = [
    (
        "доступны тарифы на 3, 6 и 12 месяцев",
        "доступны стандартные тарифы на 3 и 6 месяцев",
    ),
    (
        "Отдельно отмечу: участие в Ваших процедурах бесплатно для поставщиков, "
        "а с победителя торгов мы не удерживаем комиссию.",
        "Отдельно отмечу: участие в Ваших процедурах по стандартным тарифам — "
        "бесплатно для поставщиков и с победителя торгов мы не удерживаем комиссию.",
    ),
]


def _build_company_paragraph(paragraph) -> None:
    """Абзац о клиенте: жирное краткое название, дальше обычный текст."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(COMPANY_BOLD_PLACEHOLDER)
        paragraph.add_run(COMPANY_TAIL_PLACEHOLDER)
        return

    first = runs[0]
    first.text = COMPANY_BOLD_PLACEHOLDER
    first.bold = True
    for run in runs[1:]:
        run.text = ""

    tail = paragraph.add_run(COMPANY_TAIL_PLACEHOLDER)
    tail.bold = False
    tail.font.name = first.font.name
    tail.font.size = first.font.size
    if first.font.color and first.font.color.rgb is not None:
        tail.font.color.rgb = first.font.color.rgb


def build_kp_template() -> None:
    if not SRC_KP.exists():
        raise FileNotFoundError(f"Не найден исходный файл КП: {SRC_KP}")
    TEMPLATES.mkdir(parents=True, exist_ok=True)

    doc = Document(str(SRC_KP))
    paragraphs = doc.paragraphs

    for index, placeholder in PARAGRAPH_PLACEHOLDERS.items():
        if index >= len(paragraphs):
            raise IndexError(f"В шаблоне КП нет абзаца №{index}")
        set_paragraph_text(paragraphs[index], placeholder)

    _build_company_paragraph(paragraphs[COMPANY_PARAGRAPH])

    for old, new in INLINE_REPLACEMENTS + BLOCK_REPLACEMENTS:
        for paragraph in paragraphs:
            replace_in_paragraph(paragraph, old, new)

    doc.save(str(DST_KP))
    print(f"КП-шаблон: {DST_KP}")


def build_presentation_template() -> None:
    if not SRC_PRES.exists():
        raise FileNotFoundError(f"Не найден исходный файл презентации: {SRC_PRES}")
    TEMPLATES.mkdir(parents=True, exist_ok=True)
    # Презентация правится в приложении по тексту слайда,
    # поэтому шаблон копируется как есть — дизайн сохраняется полностью.
    shutil.copy2(SRC_PRES, DST_PRES)
    print(f"Шаблон презентации: {DST_PRES}")


def main() -> None:
    build_kp_template()
    build_presentation_template()

    # Контроль: плейсхолдеры должны присутствовать в готовом шаблоне.
    doc = Document(str(DST_KP))
    text = "\n".join(p.text for p in doc.paragraphs)
    missing = [ph for ph in PARAGRAPH_PLACEHOLDERS.values() if ph not in text]
    missing += [new for _, new in INLINE_REPLACEMENTS if new not in text]
    missing += [ph for ph in (COMPANY_BOLD_PLACEHOLDER, "{{ПОЛЬЗА_В_СЕГМЕНТЕ}}")
                if ph not in text]
    if missing:
        print("ВНИМАНИЕ, не найдены плейсхолдеры:", ", ".join(missing))
    else:
        print("Все плейсхолдеры на месте.")


if __name__ == "__main__":
    main()
